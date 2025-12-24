import sys
import os
import re
from os import path
import time
import random
import baidu_translate as fanyi
import json
from docx import Document
import time
import threading
import asyncio
#安装
#pip install python-docx


# 记录程序开始时间
start_time = time.time()
# 控制定时器是否继续执行的全局变量
stop_timer = False
# 提取数字 返回数组
def extract_numbers(s):
    # 使用正则表达式匹配浮点数或整数
    pattern = r'\d+\.\d+|\d+'
    matches = re.findall(pattern, s)
    numbers = []
    for num_str in matches:
        if '.' in num_str:
            numbers.append(float(num_str))
        else:
            numbers.append(int(num_str))
    return numbers

#翻译
async def translate_name(chinese_name):
    """简单中英翻译映射"""
    #result = fanyi.translate_text(chinese_name, to=fanyi.Lang.EN)
    result = await fanyi.translate_text_async(chinese_name, to=fanyi.Lang.EN)  # 使用异步翻译
    if not result:
        print(f"翻译失败，使用默认值：{chinese_name}")
        return ""
    else:
        # print(f"翻译成功：{chinese_name} -> {result}")
        return result

# 定义一个打印当前运行时长的函数
def print_elapsed_time():
    current_time = time.time()
    execution_time = current_time - start_time
    print(f"程序已运行 {execution_time:.2f} 秒")
    global stop_timer
    # 检查是否需要停止定时器
    if stop_timer:
        print("定时器已停止")
        return
    # 每隔10秒再次调用自己
    threading.Timer(10, print_elapsed_time).start()

# 异步写入文件的函数
async def write_to_files(notes,file_url):
    with open(file_url, 'w', encoding='utf-8') as f:
        json.dump(notes, f, ensure_ascii=False, indent=2)
    #print("数据已写入 note.json")

# 异步写入文件的函数，检查文件是否存在，不存在则创建文件并写入数据
async def write_to_file(notes, file_path):
    # 检查文件是否存在
    if not os.path.exists(file_path):
        # 如果文件不存在，创建文件并写入数据
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(notes, f, ensure_ascii=False, indent=2)
        print(f"文件 {file_path} 已创建并写入数据。")
    else:
        # 如果文件存在，则先读取现有数据，再追加新的数据
        with open(file_path, 'r+', encoding='utf-8') as f:
            try:
                existing_data = json.load(f)  # 读取现有的数据
                # 如果现有数据不是列表，初始化为空列表
                if not isinstance(existing_data, list):
                    existing_data = []
            except json.JSONDecodeError:
                # 如果文件为空或格式错误，初始化为空列表
                existing_data = []
            existing_data.extend(notes)  # 将新数据添加到现有数据
            f.seek(0)  # 将文件指针移到文件开始处
            json.dump(existing_data, f, ensure_ascii=False, indent=2)  # 写回更新的数据
        print(f"文件 {file_path} 已存在，数据已追加。")


def read_chinese_word_file():
    """
    读取中文Word文件，返回文件内容
    :param file_path: Word文件路径
    :return: 文件内容
    """
    
    # 读取Word文档
    doc = Document("chinafreq.docx")

    # 初始化一个列表，用来保存中国内地的频率划分数据
    china_mainland_data = []
    table_number = 0  # 表格编号
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历每一行
        contains_china = False  # 标记当前表格是否包含“中国内地”
        china_column_index = None  # 存储“中国内地”所在的列索引

        # 遍历表格的每一行，检查每一列
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if "中国内地" in cell.text.strip():  # 如果某列包含“中国内地”
                    contains_china = True
                    china_column_index = col_idx  # 记录该列的索引
                    break  # 找到后直接跳出列的循环
            if contains_china:  # 如果找到了“中国内地”，就跳出行的循环
                break

        # 如果当前表格不包含“中国内地”，则跳过
        if not contains_china:
            continue
        
        # 如果包含“中国内地”，则处理该表格
        print("找到包含“中国内地”的表格",china_column_index)
        current_freq_range = None
        current_additional_freq = None
        table_data = []  # 用于存储当前表格的所有数据
        table_number += 1
        # 遍历当前表格的每一行
        for row in table.rows:
            # 提取每一列中的内容
            
            cells = row.cells  # 只提取“中国内地”及其后面的列
            
            # 假设中国内地的数据在第二列
            if len(cells) >= 4:  # 确保这一行有足够的列
                china_mainland_cell = cells[china_column_index].text.strip()  # 中国内地的频率划分数据
                
                cells_items=china_mainland_cell.split("\n")

                # 处理单元格多行内容
                if len(cells_items) > 1:
                    # 如果单元格内容有多行，处理每一行
                    cells_length = len(cells_items)
                    print("单元格内容",china_mainland_cell,cells_items)
                    businesses = []
                    # 处理频率范围和附加频率
                    for cell_item_index,cell_item in enumerate(cells_items):
                        # 处理频率范围和附加频率
                        
                        if cell_item_index == 0:
                            # 处理频率范围
                            current_freq_range=extract_numbers(cell_item)
                            print("单元格频率范围",current_freq_range)
                        else:
                            current_additional_freq = cell_item.strip()
                            businesses.append(current_additional_freq)
                            print("单元格业务",current_additional_freq)
                    
                    item_data={
                            "start_freq": current_freq_range[0],
                            "end_freq": current_freq_range[1],
                            "foot_notes": [cells_items[cells_length-1].strip()],
                            "business": businesses
                        },
                    #business_info = china_mainland_cell.  # 业务名称
                    china_mainland_data.append(item_data)
        # if table_number>3:
        #     break  # 如果只需要处理前3个表格，可以在这里跳出循环








    # 将提取的数据保存为 JSON 文件
    output_data = {
        "china_mainland_frequency_division": china_mainland_data
    }

    # 定义JSON文件的路径
    json_file_path = "china_mainland_frequency_division.json"

    # 写入数据到 JSON 文件
    with open(json_file_path, "w", encoding="utf-8") as json_file:
        json.dump(output_data, json_file, ensure_ascii=False, indent=4)

    print(f"数据已成功保存为 JSON 文件：{json_file_path}")


async def read_annotation_file():
    # 启动定时器
    print_elapsed_time()
    doc=Document("annotation.docx")
    # 初始化一个列表，用来保存中国内地的频率划分数据
    # 存储结果的列表
    notes = []
    # 定义正则表达式来匹配段落中的编号（如 5.53、8.3 等）
    #note_id_pattern = re.compile(r'(\d+\.\d+)')
    # 定义正则表达式来匹配段落中的编号（如 5.53、8.3、5.54A 或其他符号后缀）
    note_id_pattern = re.compile(r'(\d+\.\d+[A-Za-z0-9\+\-\*]*)')

    read_num=0
    # 遍历文档中的所有段落
    for para in doc.paragraphs:
        # 查找段落中的编号
        read_num+=1
        match = note_id_pattern.match(para.text.strip())
        # print("段落内容",para.text)
        if match:
                # 提取编号
                note_id = match.group(1)
                # 获取编号后的文本作为语言内容
                chnese_text = para.text.strip()[len(note_id):].strip()
                en_text =await translate_name(chnese_text)
                #print("注释内容",language)
                # 创建字典并添加到结果列表
                language_cont=[
                    {
                        "code": 2052,
                        "content": chnese_text
                    },
                    {
                        "code": 1033,
                        "content": en_text
                    }
                ]
                new_note = {
                    "note_id": note_id,
                    "language": language_cont
                }
                notes.append(new_note)
                # 异步写入文件
                output_data = {
                    "notes_list": notes
                }
                await write_to_files(output_data,"notes_list.json")
        # if read_num>10:
        #     print("读取注释文档完成",notes)
        #     break
        
    # 将提取的数据保存为 JSON 文件
    output_data = {
        "notes_list": notes
    }
    await write_to_files(output_data,"notes_list.json")
    # 定义JSON文件的路径
    json_file_path = "notes_list.json"
    # 最终写入文件
    # asyncio.run(write_to_file(output_data,"notes_list.json"))
    # 写入数据到 JSON 文件
    # with open(json_file_path, "w", encoding="utf-8") as json_file:
    #     json.dump(output_data, json_file, ensure_ascii=False, indent=4)
    global stop_timer
    stop_timer = True  # 停止定时器
    print(f"数据已成功保存为 JSON 文件")
        

async def read_annotation_file3(out_file):
    # 启动定时器
    print_elapsed_time()
    doc=Document("annotation.docx")
    # 初始化一个列表，用来保存中国内地的频率划分数据
    # 存储结果的列表
    notes = []

    # 优化的正则表达式：
    #     \d+\.\d+：匹配数字和小数点，确保编号部分正确。
    #     [A-Za-z0-9\+\-\*]*：匹配编号后可能出现的字母、数字或其他符号。
    #     .*：匹配后面内容，包括空格、中文、标点符号等。
    # 解释：
    #     ^: 确保匹配从段落的开头。
    #     (\d+\.\d+[A-Za-z0-9\+\-\*]*): 匹配编号部分。
    #     \s+: 匹配编号后的空格。
    #     (.*): 捕获编号后的所有内容作为 content。
    note_id_pattern = re.compile(r'^(\d+\.\d+[A-Za-z0-9\+\-\*]*)\s+(.*)')

    read_num=0
    isread_text=False
    # 遍历文档中的所有段落
    for para in doc.paragraphs:
        # 查找段落中的编号
        read_num+=1
        match = note_id_pattern.match(para.text.strip())
        # print("段落内容",para.text)
        if match:
            note_ids = match.group(1)
            # 获取编号后的文本作为语言内容
            if note_ids=="3.6":
                isread_text=True
                print("开始读取注释内容",para.text.strip())


        if isread_text:
                # 提取编号
                # 分割每段文字，假设每段内容按“编号 空格 内容”结构
                parts = para.text.split(" ", 1)
                print("开始读取注释以后内容",para.text.strip())
                if len(parts) == 2:
                    print("提取提取",parts[0].strip())
                    note_id = parts[0].strip()  # 提取编号部分
                    # 获取编号后的文本作为语言内容
                    chnese_text = parts[1].strip()
                    en_text =await translate_name(chnese_text)
                    #print("注释内容",language)
                    # 创建字典并添加到结果列表
                    language_cont=[
                        {
                            "code": 2052,
                            "content": chnese_text
                        },
                        {
                            "code": 1033,
                            "content": en_text
                        }
                    ]
                    new_note = {
                        "note_id": note_id,
                        "language": language_cont
                    }
                    notes.append(new_note)
                    # 异步写入文件
                    output_data = {
                        "notes_list": notes
                    }
                    await write_to_files(output_data,out_file)
        # if read_num>10:
        #     print("读取注释文档完成",notes)
        #     break
        
    # 将提取的数据保存为 JSON 文件
    output_data = {
        "notes_list": notes
    }
    await write_to_files(output_data,out_file)
    global stop_timer
    stop_timer = True  # 停止定时器
    print(f"数据已成功保存为 JSON 文件")
        

# 启动异步任务
async def main():
    # 读取Word文件并提取数据
    read_chinese_word_file()
    # 读取注释文档
    start_time = time.time()
    #note=await read_annotation_file()
    # note=await read_annotation_file3("notes_list3.json")
    end_time = time.time()
    print(f"读取注释文档完成，耗时：{end_time - start_time}秒")

    

# 使用 asyncio.run() 来启动程序
if __name__ == "__main__":
     loop = asyncio.get_event_loop()
     loop.run_until_complete(main())
        