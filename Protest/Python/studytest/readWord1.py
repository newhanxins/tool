import sys
import os
import re
from os import path
import time
import random
import baidu_translate as fanyi
import json
from docx import Document
import time

#安装
#pip install python-docx

# 提取数字 返回数组
def extract_numbers(s):
    # 使用正则表达式匹配浮点数或整数
    pattern = r'\d+\.\d+|\d+'
    matches = re.findall(pattern, s)
    numbers = []
    for num_str in matches:
        if '.' in num_str:
            numbers.append(float(num_str))
        else:
            numbers.append(int(num_str))
    return numbers

#翻译
def translate_name(chinese_name):
    """简单中英翻译映射"""
    result = fanyi.translate_text(chinese_name, to=fanyi.Lang.EN)
    if not result:
        print(f"翻译失败，使用默认值：{chinese_name}")
        return ""
    else:
        print(f"翻译成功：{chinese_name} -> {result}")
        return result

def read_chinese_word_file():
    """
    读取中文Word文件，返回文件内容
    :param file_path: Word文件路径
    :return: 文件内容
    """
    
    # 读取Word文档
    doc = Document("chinafreq.docx")

    # 初始化一个列表，用来保存中国内地的频率划分数据
    china_mainland_data = []
    table_number = 0  # 表格编号
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历每一行
        contains_china = False  # 标记当前表格是否包含“中国内地”
        china_column_index = None  # 存储“中国内地”所在的列索引

        # 遍历表格的每一行，检查每一列
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if "中国内地" in cell.text.strip():  # 如果某列包含“中国内地”
                    contains_china = True
                    china_column_index = col_idx  # 记录该列的索引
                    break  # 找到后直接跳出列的循环
            if contains_china:  # 如果找到了“中国内地”，就跳出行的循环
                break

        # 如果当前表格不包含“中国内地”，则跳过
        if not contains_china:
            continue
        
        # 如果包含“中国内地”，则处理该表格
        print("找到包含“中国内地”的表格",china_column_index)
        current_freq_range = None
        current_additional_freq = None
        table_data = []  # 用于存储当前表格的所有数据
        table_number += 1
        # 遍历当前表格的每一行
        for row in table.rows:
            # 提取每一列中的内容
            
            cells = row.cells  # 只提取“中国内地”及其后面的列
            
            # 假设中国内地的数据在第二列
            if len(cells) >= 4:  # 确保这一行有足够的列
                china_mainland_cell = cells[china_column_index].text.strip()  # 中国内地的频率划分数据
                
                cells_items=china_mainland_cell.split("\n")

                # 处理单元格多行内容
                if len(cells_items) > 1:
                    # 如果单元格内容有多行，处理每一行
                    cells_length = len(cells_items)
                    print("单元格内容",china_mainland_cell,cells_items)
                    businesses = []
                    # 处理频率范围和附加频率
                    for cell_item_index,cell_item in enumerate(cells_items):
                        # 处理频率范围和附加频率
                        
                        if cell_item_index == 0:
                            # 处理频率范围
                            current_freq_range=extract_numbers(cell_item)
                            print("单元格频率范围",current_freq_range)
                        else:
                            current_additional_freq = cell_item.strip()
                            businesses.append(current_additional_freq)
                            print("单元格业务",current_additional_freq)
                    
                    item_data={
                            "start_freq": current_freq_range[0],
                            "end_freq": current_freq_range[1],
                            "foot_notes": [cells_items[cells_length-1].strip()],
                            "business": businesses
                        },
                    #business_info = china_mainland_cell.  # 业务名称
                    china_mainland_data.append(item_data)
        # if table_number>3:
        #     break  # 如果只需要处理前3个表格，可以在这里跳出循环








    # 将提取的数据保存为 JSON 文件
    output_data = {
        "china_mainland_frequency_division": china_mainland_data
    }

    # 定义JSON文件的路径
    json_file_path = "china_mainland_frequency_division.json"

    # 写入数据到 JSON 文件
    with open(json_file_path, "w", encoding="utf-8") as json_file:
        json.dump(output_data, json_file, ensure_ascii=False, indent=4)

    print(f"数据已成功保存为 JSON 文件：{json_file_path}")


def read_annotation_file():
    doc=Document("annotation.docx")
    # 初始化一个列表，用来保存中国内地的频率划分数据
    # 存储结果的列表
    notes = []
    # 定义正则表达式来匹配段落中的编号（如 5.53、8.3 等）
    #note_id_pattern = re.compile(r'(\d+\.\d+)')
    # 定义正则表达式来匹配段落中的编号（如 5.53、8.3、5.54A 或其他符号后缀）
    note_id_pattern = re.compile(r'(\d+\.\d+[A-Za-z0-9\+\-\*]*)')

    read_num=0
    # 遍历文档中的所有段落
    for para in doc.paragraphs:
        # 查找段落中的编号
        read_num+=1
        match = note_id_pattern.match(para.text.strip())
        # print("段落内容",para.text)
        if match:
                # 提取编号
                note_id = match.group(1)
                # 获取编号后的文本作为语言内容
                chnese_text = para.text.strip()[len(note_id):].strip()
                en_text =translate_name(chnese_text)
                #print("注释内容",language)
                # 创建字典并添加到结果列表
                language_cont=[
                    {
                        "code": 2052,
                        "content": chnese_text
                    },
                    {
                        "code": 1033,
                        "content": en_text
                    }
                ]
                notes.append({
                    "note_id": note_id,
                    "language": language_cont
                })
        if read_num>10:
            print("读取注释文档完成",notes)
            break
        
    # 将提取的数据保存为 JSON 文件
    output_data = {
        "notes_list": notes
    }

    # 定义JSON文件的路径
    json_file_path = "notes_list.json"

    # 写入数据到 JSON 文件
    with open(json_file_path, "w", encoding="utf-8") as json_file:
        json.dump(output_data, json_file, ensure_ascii=False, indent=4)

    print(f"数据已成功保存为 JSON 文件：{output_data}")
        

def main():
    # 读取Word文件并提取数据
    #read_chinese_word_file()
    # 读取注释文档
    start_time = time.time()
    read_annotation_file()
    end_time = time.time()
    print(f"读取注释文档完成，耗时：{end_time - start_time}秒")

    

if __name__ == "__main__":
    main()